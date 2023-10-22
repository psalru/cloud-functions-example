import os
import re
import json
import docx
import pandas as pd
from docx.shared import Mm
from datetime import datetime
import matplotlib.pyplot as plt


def gen_report(settings: dict):
    university_id = settings['id']
    data_folder = f"{settings['folder']}/{university_id}"
    doc = docx.Document('template.docx')
    data = pd.read_csv(f"{data_folder}/data.csv", index_col=0)
    tables = {}

    # Вставляем параметры

    for paragraph in doc.paragraphs:
        found_param = re.finditer(r"\{\{\w+\}\}", paragraph.text)

        if found_param:
            for match in found_param:
                fp = re.sub(r'[\{\}]', '', match.group())

                if fp in data.index:
                    value = data.loc[fp]['value']

                    paragraph.text = paragraph.text.replace('{{' + str(fp) + '}}', str(value) if not pd.isnull(value) else '')

    # Вставляем изображения

    for paragraph in doc.paragraphs:
        found_image = re.finditer(r"\[\[[^\]]+\]\]", paragraph.text)

        if found_image:
            for match in found_image:
                fi = re.sub(r'[\[\]]', '', match.group().strip())
                image_with_param = f"{data_folder}/{fi}".split(' ')
                image = image_with_param[0]

                if os.path.isfile(image):
                    paragraph.text = ''
                    r = paragraph.add_run()

                    if len(image_with_param) > 1:
                        pic_width = int(image_with_param[1])
                        r.add_picture(image, width=Mm(pic_width))
                    else:
                        r.add_picture(image)

    # Вставляем данные в таблицы
    # styles = doc.styles
    #
    # print([x.name for x in styles])

    # Сначала находим таблицы
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    annexes_param = re.findall(r"\(\(\w+\)\)", p.text)

                    if annexes_param:
                        annex_title = re.sub(r'[\(\)]', '', annexes_param[0])
                        p.text = p.text.replace(annexes_param[0], '').strip()

                        if annex_title not in tables.keys():
                            tables[annex_title] = table

    # Потом вставляем данные в таблицы
    for key in tables.keys():
        annex_df_file_path = f"{data_folder}/{key}.csv"

        if os.path.isfile(annex_df_file_path):
            annex_df = pd.read_csv(annex_df_file_path)
            table = tables[key]
            table.style = 'Normal Table'

            for head_i, head in annex_df.iterrows():
                cells = table.add_row().cells

                for col_i, col in enumerate(annex_df.columns):
                    cells[col_i].text = str(head[col])

    # Сохраняем отчёт

    doc.save(f"{data_folder}/report_{university_id}.docx")


def gen_content(settings: dict):
    university_id, fig_width = settings['id'], settings['width']
    data_folder = f"{settings['folder']}/{university_id}"
    texts = json.load(open('texts.json', 'r'))
    df = pd.read_csv('https://storage.yandexcloud.net/psal.public/hosts/psal/dumps/hh_university_vacancies_by_month.csv', sep='|', index_col=0)
    university = df[df['university_id'] == university_id][['university_abbreviation', 'id']].\
        groupby(by='university_abbreviation', as_index=False).count().sort_values(by='id', ascending=False)['university_abbreviation'].unique()[0]
    data = pd.DataFrame(columns=['value'])
    data.loc['datetime', 'value'] = datetime.now().strftime('%d.%m.%Y %H:%M')
    data.loc['university', 'value'] = university
    df['salary'] = df.apply(lambda x: x['salary_to'] if not pd.isnull(x['salary_to']) else x['salary_from'], axis=1)
    region = df[df['university_abbreviation'] == university][['region', 'id']].groupby(by='region', as_index=False).count().sort_values(by='id', ascending=False).iloc[0]['region']

    if not os.path.isdir(data_folder):
        os.mkdir(data_folder)

    # Данные про преподавателей

    df_teachers = df[df['professional_roles'] == 'Учитель, преподаватель, педагог'].copy()
    df_teachers_by_university = df_teachers[df_teachers['university_id'] == university_id].copy().sort_values(by='salary', ascending=False)
    teachers_mean_by_rf = df_teachers['salary'].mean() / 1000
    teachers_mean_by_region = df_teachers[df_teachers['region'] == region]['salary'].mean() / 1000
    stat_by_teachers = pd.DataFrame([
        {'title': 'РФ', 'value': teachers_mean_by_rf},
        {'title': region, 'value': teachers_mean_by_region}
    ]).dropna()

    if len(df_teachers_by_university) > 0:
        teachers_mean_by_university = df_teachers_by_university['salary'].mean() / 1000
        teachers_top_by_university = df_teachers_by_university.sort_values(by='salary', ascending=False).iloc[0]
        data.loc['description_teachers', 'value'] = texts['not_empty'].format(
            prof_role='Учитель, преподаватель, педагог',
            percent_of_rf='{0:0.2%}'.format(teachers_mean_by_university / teachers_mean_by_rf),
            percent_of_region='{0:0.2%}'.format(teachers_mean_by_university / teachers_mean_by_region),
            top_salary=round(teachers_top_by_university['salary'] / 1000, 0),
            top_vacancy=teachers_top_by_university['title'],
            top_url=teachers_top_by_university['url']
        )
        stat_by_teachers.loc[len(stat_by_teachers)] = {
            'title': university,
            'value': teachers_mean_by_university
        }
    else:
        data.loc['description_teachers', 'value'] = texts['empty'].format(
            prof_role='Учитель, преподаватель, педагог'
        )
        stat_by_teachers.loc[len(stat_by_teachers)] = {
            'title': university,
            'value': 0
        }

    fig, ax = plt.subplots()
    ax.bar(stat_by_teachers['title'], stat_by_teachers['value'], color='green')
    ax.set_ylim(0, stat_by_teachers['value'].max() * 1.15)

    for i, r in stat_by_teachers.iterrows():
        ax.annotate(
            '{0:.2f} тыс. ₽'.format(r['value']),
            (r['title'], r['value']),
            va='bottom', ha='center', xytext=(0, 10),
            textcoords='offset points', fontweight='bold',
        )

    fig.set_figwidth(fig_width)
    fig.savefig(f"{data_folder}/stat_by_teachers.png", format='png', bbox_inches='tight')
    plt.close(fig)

    # Про исследователей

    df_researcher = df[df['professional_roles'] == 'Научный специалист, исследователь'].copy()
    df_researcher_by_university = df_researcher[df_researcher['university_id'] == university_id].copy().sort_values(by='salary', ascending=False)
    researchers_mean_by_rf = df_researcher['salary'].mean() / 1000
    researchers_mean_by_region = df_researcher[df_researcher['region'] == region]['salary'].mean() / 1000
    stat_by_researcher = pd.DataFrame([
        {'title': 'РФ', 'value': df_researcher['salary'].mean() / 1000},
        {'title': region, 'value': df_researcher[df_researcher['region'] == region]['salary'].mean() / 1000},
    ]).dropna()

    if len(df_researcher_by_university) > 0:
        researchers_mean_by_university = df_researcher_by_university['salary'].mean() / 1000
        researchers_top_by_university = df_researcher_by_university.iloc[0]
        data.loc['description_researcher', 'value'] = texts['not_empty'].format(
            prof_role='Научный специалист, исследователь',
            percent_of_rf='{0:0.2%}'.format(researchers_mean_by_university / researchers_mean_by_rf),
            percent_of_region='{0:0.2%}'.format(researchers_mean_by_university / researchers_mean_by_region),
            top_salary=round(researchers_top_by_university['salary'] / 1000, 0),
            top_vacancy=researchers_top_by_university['title'],
            top_url=researchers_top_by_university['url']
        )
        stat_by_researcher.loc[len(stat_by_researcher)] = {
          'title': university,
          'value': researchers_mean_by_university
        }
    else:
        data.loc['description_researcher', 'value'] = texts['empty'].format(prof_role='Учитель, преподаватель, педагог')
        stat_by_researcher.loc[len(stat_by_researcher)] = {
            'title': university,
            'value': 0
        }

    fig, ax = plt.subplots()
    ax.bar(stat_by_researcher['title'], stat_by_researcher['value'], color='orange')
    ax.set_ylim(0, stat_by_researcher['value'].max() * 1.15)

    for i, r in stat_by_researcher.iterrows():
        ax.annotate(
            '{0:.2f} тыс. ₽'.format(r['value']),
            (r['title'], r['value']),
            va='bottom', ha='center', xytext=(0, 10),
            textcoords='offset points', fontweight='bold',
        )

    fig.set_figwidth(fig_width)
    fig.savefig(f"{data_folder}/stat_by_researcher.png", format='png', bbox_inches='tight')
    plt.close(fig)

    # Распределение вакансий по проф. ролям

    teachers_count = len(df_teachers[df_teachers['university_id'] == university_id])
    researchers_count = len(df_researcher[df_researcher['university_id'] == university_id])
    stat_by_professional_roles = pd.DataFrame([
        {
            'title': 'ППС',
            'value': teachers_count,
            'percent': teachers_count / len(df_teachers)
        },
        {
            'title': 'НР',
            'value': researchers_count,
            'percent': researchers_count / len(df_researcher)
        },
        {
            'title': 'Остальное',
            'value': len(df[df['university_id'] == university_id]) - teachers_count - researchers_count,
            'percent': (len(df[df['university_id'] == university_id]) - teachers_count - researchers_count) / len(df)
        },
    ])

    fig, ax = plt.subplots()
    ax.bar(stat_by_professional_roles['title'], stat_by_professional_roles['value'], color=['green', 'orange', 'gray'])
    ax.set_ylim(0, stat_by_professional_roles['value'].max() * 1.15)

    for i, r in stat_by_professional_roles.iterrows():
        ax.annotate(
            '{0:.0f} ({1:.2%}*)'.format(r['value'], r['percent']),
            (r['title'], r['value']),
            va='bottom', ha='center', xytext=(0, 10),
            textcoords='offset points', fontweight='bold',
        )

    fig.set_figwidth(fig_width)
    fig.savefig(f"{data_folder}/stat_by_professional_roles.png", format='png', bbox_inches='tight')
    plt.close(fig)

    # ТОП 10 вакансий по заработной плате

    df_top_by_salary = df[~df['salary'].isnull()].sort_values(by='salary', ascending=False).reset_index()[:10]
    annex_top_by_salary = df_top_by_salary[['title', 'salary', 'university_abbreviation', 'url']].copy()
    annex_top_by_salary['salary'] = annex_top_by_salary['salary'].astype(int)
    data.loc['top_vacancies', 'value'] = texts['top_by_salary'].format(
        regions='; '.join([f"{i} ({x['id']})" for i, x in df_top_by_salary[['region', 'id']].groupby(by='region').count().sort_values(by='id', ascending=False).iterrows()]),
        count=len(df_top_by_salary['professional_roles'].unique()),
        prof_roles='; '.join([f"{i} ({x['id']})" for i, x in df_top_by_salary[['professional_roles', 'id']].groupby(by='professional_roles').count().sort_values(by='id', ascending=False).iterrows()])
    )

    # Сохраняем результаты

    data.to_csv(f"{data_folder}/data.csv")
    annex_top_by_salary.to_csv(f"{data_folder}/annex_top_by_salary.csv", index=False)


